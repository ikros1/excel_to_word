from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import serialize_for_reading
from docx.oxml.ns import nsmap, qn

def read_word_document(file_path):
    doc = Document(file_path)
    paragraphs = [paragraph for paragraph in doc.paragraphs]
    tables = [table for table in doc.tables]
    return paragraphs, tables

def get_font_info(run):
    return {
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': run.font.color.rgb if run.font.color.rgb is not None else None,
        'size': run.font.size if run.font.size is not None else Pt(12),
        'name': run.font.name,
    }

def get_paragraph_font_info(paragraph):
    p_rpr = paragraph.style.element.xpath('w:rPr')[0]
    #print(p_rpr[0].xml)
    if p_rpr.xpath('w:rFonts'):
        try:
            return p_rpr.xpath('w:rFonts')[0].attrib[qn("w:eastAsia")]
        except:
            return p_rpr.xpath('w:rFonts')[0].attrib[qn("w:ascii")]
    return None

def copy_characters(source_paragraphs, target_doc):
    for source_paragraph in source_paragraphs:
        new_paragraph = target_doc.add_paragraph()

        # 复制段落格式
        new_paragraph._element.clear_content()

        # 获取段落字体信息
        paragraph_font = get_paragraph_font_info(source_paragraph)

        for run in source_paragraph.runs:
            new_run = new_paragraph.add_run(run.text)

            # 输出每个字符的字体信息
            for char in run.text:
                print(f"Character: {char}")

                # 获取字符字体信息
                font_info = get_font_info(run)

                # 如果源字体信息为 None，使用段落字体信息
                if font_info['name'] is None:
                    font_info['name'] = paragraph_font

                print(f"Font Bold: {font_info['bold']}")
                print(f"Font Italic: {font_info['italic']}")
                print(f"Font Underline: {font_info['underline']}")
                print(f"Font Color RGB: {font_info['color']}")
                print(f"Font Size: {font_info['size']}")
                print(f"Font Name: {font_info['name']}")
                print("---")

                # 复制字符格式
                new_run.bold = font_info['bold']
                new_run.italic = font_info['italic']
                new_run.underline = font_info['underline']
                new_run.font.color.rgb = font_info['color']
                new_run.font.size = font_info['size']
                new_run.font.name = font_info['name']

        # 复制段落属性
        new_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
        new_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
        new_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
        new_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
        new_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing

# 用法示例
input_file_path = 'word_file/example.docx'
paragraphs, tables = read_word_document(input_file_path)

# 创建一个新的 Document 对象
doc = Document()

# 添加标题
doc.add_heading('My Document Title', level=1)

# 复制字符
copy_characters(paragraphs, doc)

# 添加带有项目符号的列表
doc.add_paragraph('This is a bulleted list:', style='ListBullet')
doc.add_paragraph('Item 1')
doc.add_paragraph('Item 2')

# 保存.docx文件
output_file_path = 'word_file/write_file.docx'
doc.save(output_file_path)

print('Document created successfully!')
