import threading
import time
import openpyxl
from docx import Document
import os
import shutil
import zipfile


class TaskInstance:
    # 用户实例执行单元任务类
    # 实现对用户上传接收，初始化任务，检查任务， 执行任务， 返回任务结果， 返回费用详情， 查询任务进度与状态，超时和已完成任务资源的回收的实现
    def __init__(self, index, from_person_id, passwd):
        # 执行密码
        self.passwd = str(passwd)
        # 任务队列编号
        self.index = index
        # 任务来自队员编号
        self.from_person_id = from_person_id
        # 任务当前状态
        self.status = "check"
        # all_status is
        # check
        # ready
        # start
        # working
        # error
        # ready_to_back
        # out_time
        # need_to_delete
        # checked_error
        # 任务当前状态信息描述
        self.status_info = ""
        # 需要支付的金额
        self.pay = 0
        # 任务开始的时间
        self.start_time = 0
        # 任务结束的时间
        self.over_time = 0
        # 任务word文件名
        self.word_file = ""
        # 任务excel文件名
        self.excel_file = ""
        # 上传excel文件路径
        self.upload_excel_file_path = "uploads/excel/"
        # 上传word文件路径
        self.upload_word_file_path = "uploads/word/"
        # 发送文件路径
        self.send_file_path = "output/" + str(self.from_person_id) + "/"
        os.makedirs(self.send_file_path, exist_ok=True)
        # 用户指定名格式种类  #写不完后期写
        self.user_like_words = ""
        # 用户任务完成百分比
        self.task_over_percent = 0
        # 任务初始化进行中标志
        self.task_init_ing_flag = True
        # 替换词列表
        self.replacements_list = []
        # 返回文件包名
        self.send_zip_file_path_name = ""
        # 读取实例化的word和excel对应的handle
        self.word_file_handle = None
        self.excel_file_handle = None
        self.last_one_preview_file_path_name = ""
        self.magnification = 1000

    def make_replace_list(self):
        print("start make_replace_list from id: " + self.from_person_id)
        try:
            self.start_time = time.time()
            self.excel_file_handle = openpyxl.load_workbook(self.upload_excel_file_path + self.excel_file)
            sheet = self.excel_file_handle.active
            # First loop: Find columns with "er_wr_" and corresponding elements
            columns_with_marker = {}
            for col_idx in range(1, sheet.max_column + 1):
                for row_idx in range(1, sheet.max_row + 1):
                    cell_value = sheet.cell(row=row_idx, column=col_idx).value
                    if cell_value is not None and "er_wr_" in str(cell_value):
                        if col_idx not in columns_with_marker:
                            columns_with_marker[col_idx] = []
                        columns_with_marker[col_idx].append((row_idx, cell_value))

            # Second loop: Build the return_list
            for row_idx in range(1, sheet.max_row + 1):
                row_data = []
                for col_idx, markers in columns_with_marker.items():
                    if markers:
                        cell_value = sheet.cell(row=row_idx, column=col_idx).value
                        if cell_value is not None:
                            row_data.append((cell_value, markers[0][1]))
                if row_data:
                    self.replacements_list.append(row_data)
            print("make_replace_list success from id: " + self.from_person_id)
        except Exception as e:
            self.status_info = f"Error in Task {self.index} make replace list error: {str(e)} "
            self.status = "error"
            print("make replace error from id : " + self.from_person_id + "error: " + str(e))

    def set_word_file(self, word_file_name):
        self.word_file = word_file_name
        print("set word file: " + self.word_file + "from id: " + self.from_person_id)

    def set_excel_file(self, excel_file_name):
        self.excel_file = excel_file_name
        print("set excel file: " + self.excel_file + "from id: " + self.from_person_id)

    def check_all_file(self):
        print("start check all file from id : " + self.from_person_id)
        self.status_info = f"Task {self.index} from {self.from_person_id} checking"
        self.status = "check"
        if self.word_file == "" or self.excel_file == "":
            self.status = "error"
            self.status_info = "not exist word_file and excel_file"
        else:
            try:
                self.excel_file_handle = openpyxl.load_workbook(self.upload_excel_file_path + self.excel_file)
                self.word_file_handle = Document(self.upload_word_file_path + self.word_file)
                self.status = "ready"
                self.status_info = "all_ready"
                self.task_init_ing_flag = False
                print("check all file success from id : " + self.from_person_id)
            except Exception as e:
                self.status_info = f"Error in Task {self.index}: {str(e)} check excel or word file error"
                self.status = "error"
                print("check all file error from id : " + self.from_person_id + "error: " + str(e))

    def replace_text_in_docx(self, save_file_name, replacements):
        self.word_file_handle = Document(self.upload_word_file_path + self.word_file)
        for para in self.word_file_handle.paragraphs:
            for run in para.runs:
                for new_text, old_text in replacements:
                    # 类型转换string
                    old_text = str(old_text)
                    new_text = str(new_text)
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

        self.last_one_preview_file_path_name = self.send_file_path + save_file_name
        self.word_file_handle.save(self.send_file_path + save_file_name)
        print("replace text in docx success from id : " + self.from_person_id + "save file name :" + save_file_name)

    def delete_output_file(self):
        print("start delete all file from id : " + self.from_person_id)
        try:
            os.remove(self.send_zip_file_path_name)
            print(f"Deleted file: {self.send_zip_file_path_name}")
        except Exception as e:
            print(f"Error deleting file {self.send_zip_file_path_name}: {e}")
        print("start delete all file from id : " + self.from_person_id)
        try:
            os.remove(self.upload_excel_file_path+self.word_file)
            print(f"Deleted file: {self.upload_excel_file_path+self.word_file}")
        except Exception as e:
            print(f"Error deleting file {self.upload_excel_file_path+self.word_file}: {e}")
        try:
            os.remove(self.upload_excel_file_path+self.excel_file)
            print(f"Deleted file: {self.upload_excel_file_path+self.excel_file}")
        except Exception as e:
            print(f"Error deleting file {self.upload_excel_file_path+self.excel_file}: {e}")
        for root, dirs, files in os.walk(self.send_file_path):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    os.remove(file_path)
                    print(f"Deleted file: {file_path}")
                except Exception as e:
                    print(f"Error deleting file {file_path}: {e}")

            for dir in dirs:
                dir_path = os.path.join(root, dir)
                try:
                    shutil.rmtree(dir_path)
                    print(f"Deleted folder: {dir_path}")
                except Exception as e:
                    print(f"Error deleting folder {dir_path}: {e}")

    def zip_folder(self):
        if self.status == "ready_to_back":
            folder_path = self.send_file_path
            zip_path = self.send_file_path + str(self.from_person_id) + ".zip"
            self.send_zip_file_path_name = self.send_file_path + str(self.from_person_id) + ".zip"
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_f:
                for root, dirs, files in os.walk(folder_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, folder_path)
                        zip_f.write(file_path, arcname=arc_name)

    def all_file_replace_text_in_docx(self):
        print("start all file replace text in docx from id : " + self.from_person_id)
        for i, replacements in enumerate(self.replacements_list):
            try:
                # 个性化处理命名方式实现位置，没写，不知道咋写
                save_file_name = self.user_like_words + str(i) + ".doc"
                # 任务百分比进程实现
                decimal_number = i+1 / len(self.replacements_list)
                if decimal_number == 1:
                    self.over_time = time.time()
                    self.pay = (self.over_time - self.start_time) * self.magnification
                    self.status_info = f"Task {self.index}: done"
                    self.status = "ready_to_back"
                rounded_number = round(decimal_number, 2)
                # 将保留小数后的数字转换为字符串
                self.task_over_percent = str(rounded_number)
                print(
                    "file replace text in docx from id : " + self.from_person_id + " percent: " + self.task_over_percent + " have done")
                self.replace_text_in_docx(save_file_name, replacements)
            except Exception as e:
                self.status_info = f"Error in Task {self.index}: {str(e)}"
                self.status = "error"
                print("start all file replace text in docx error from id : " + self.from_person_id + "error: " + str(e))

    def execute(self):
        try:
            # 在这里执行实际的任务逻辑
            time.sleep(5)
            self.check_all_file()
            if self.task_init_ing_flag:
                time.sleep(5)
                self.make_replace_list()
                if self.status != "error":
                    self.all_file_replace_text_in_docx()
                if self.status != "error":
                    self.zip_folder()
        except Exception as e:
            self.status_info = f"Error in Task {self.index}: {str(e)}"
            self.status = "error"

    def check_out_time(self):
        while True:
            if self.status == "error":
                self.over_time = time.time()
            start_time_long = time.time() - self.start_time
            over_time_long = time.time() - self.over_time
            if start_time_long > 5000 or over_time_long > 500 or self.status == "need_to_delete":
                self.status_info = f"delete Task {str(self.index)} from id : {str(self.from_person_id)}"
                self.status = "delete_ing"
                print(self.status_info)
                self.delete_output_file()
                self.status = "deleted"

    def get_status(self):
        now = self.status
        if self.status == "error":
            self.status = "need_to_delete"

        return now

    def set_have_get_file_status(self):
        self.status = "need_to_delete"

    def start(self):
        thread = threading.Thread(target=self.execute)
        thread.start()
        thread = threading.Thread(target=self.check_out_time)
        thread.start()
