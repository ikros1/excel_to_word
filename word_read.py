from docx import Document

def read_word_document(file_path):
    # 打开Word文档
    doc = Document(file_path)

    # 读取文档的每个段落
    for paragraph in doc.paragraphs:
        print(paragraph.text)

    # 读取文档的每个表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

# 用法示例
file_path = 'word_file/example.docx'
read_word_document(file_path)
