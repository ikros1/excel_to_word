from docx import Document

def find_text_in_word(document_path, target_text):
    doc = Document(document_path)

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        if target_text in paragraph.text:
            print(f"Found '{target_text}' in paragraph {paragraph_index + 1}")

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if target_text in cell.text:
                    print(f"Found '{target_text}' in table cell ({table_index + 1}, {row_index + 1}, {cell_index + 1})")

if __name__ == "__main__":
    document_path = "word_file/read_file.docx"
    target_text = "截图"

    find_text_in_word(document_path, target_text)
