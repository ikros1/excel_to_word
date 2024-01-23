from xml.dom import minidom
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def read_word_document(file_path, xml_file_path, input_font_path):
    doc_f = Document(file_path)
    paragraphs_f = [paragraph for paragraph in doc_f.paragraphs]
    tables_f = [table for table in doc_f.tables]
    return paragraphs_f, tables_f, xml_file_path, input_font_path


def get_font_info(run):
    return {
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': run.font.color.rgb if run.font.color.rgb is not None else None,
        'size': run.font.size if run.font.size is not None else Pt(12),
        'name': run.font.name,
    }


def get_defaut_font_info_from_fontxml(file_path, find_style):
    # 解析XML文件
    dom = minidom.parse(file_path)

    # 获取<w:styles>元素
    styles_elements = dom.getElementsByTagName('w:font')
    if styles_elements:
        for styles_element in styles_elements:
            if find_style == styles_element.getAttribute('w:name'):
                return styles_element.getElementsByTagName('w:altName')[0].getAttribute('w:val')
    else:
        return None


def get_paragraph_font_info(style_file_path, font_file_path):
    # 解析XML文件
    dom = minidom.parse(style_file_path)

    # 获取<w:styles>元素
    styles_element = dom.getElementsByTagName('w:styles')
    if styles_element:
        styles_element = styles_element[0]

        # 获取<w:rFonts>元素
        r_fonts_element = styles_element.getElementsByTagName('w:rFonts')
        if r_fonts_element:
            r_fonts_element = r_fonts_element[0]

            # 输出属性值
            print("Attribute: w:ascii =", r_fonts_element.getAttribute('w:ascii'))
            print("Attribute: w:eastAsia =", r_fonts_element.getAttribute('w:eastAsia'))
            print("Attribute: w:hAnsi =", r_fonts_element.getAttribute('w:hAnsi'))
            print("Attribute: w:cs =", r_fonts_element.getAttribute('w:cs'))
            if r_fonts_element is not None:
                r_fonts_values = {
                    'ascii': r_fonts_element.getAttribute('w:ascii'),
                    'eastAsia': r_fonts_element.getAttribute('w:eastAsia'),
                    'hAnsi': r_fonts_element.getAttribute('w:hAnsi'),
                    'cs': r_fonts_element.getAttribute('w:cs')
                }
                if r_fonts_values['eastAsia'] is not None:
                    turn_to_style = r_fonts_values['eastAsia']
                elif r_fonts_values['cs'] is not None:
                    turn_to_style = r_fonts_values['cs']
                elif r_fonts_values['ascii'] is not None:
                    turn_to_style = r_fonts_values['ascii']
                else:
                    turn_to_style = r_fonts_values['hAnsi']
                return get_defaut_font_info_from_fontxml(font_xml_file, turn_to_style)
            else:
                return None

        else:
            return None
    else:
        print("No <w:docDefaults> element found in the XML.")


def copy_characters(source_paragraphs, target_doc, xml_file, font_xml_file):
    for source_paragraph in source_paragraphs:
        new_paragraph = target_doc.add_paragraph()

        # 复制段落格式
        new_paragraph._element.clear_content()

        # 获取段落字体信息
        paragraph_font = get_paragraph_font_info(xml_file, font_xml_file)

        for run in source_paragraph.runs:
            new_run = new_paragraph.add_run(run.text)

            # 输出每个字符的字体信息
            for char in run.text:
                print(f"Character: {char}")

                # 获取字符字体信息
                font_info = get_font_info(run)

                # 如果源字体信息为 None，使用段落字体信息
                if font_info['name'] is None:
                    font_info['name'] = paragraph_font

                print(f"Font Bold: {font_info['bold']}")
                print(f"Font Italic: {font_info['italic']}")
                print(f"Font Underline: {font_info['underline']}")
                print(f"Font Color RGB: {font_info['color']}")
                print(f"Font Size: {font_info['size']}")
                print(f"Font Name: {font_info['name']}")
                print("---")

                # 复制字符格式
                new_run.font.name = font_info['name']
                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_info['name'])

                new_run.bold = font_info['bold']
                new_run.italic = font_info['italic']
                new_run.underline = font_info['underline']
                new_run.font.color.rgb = font_info['color']
                new_run.font.size = font_info['size']


        # 复制段落属性
        new_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
        new_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
        new_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
        new_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
        new_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing


# 用法示例
input_file_path = 'word_file/example.docx'
input_file_xml_path = "word_file/example_zip/word/styles.xml"
input_font_xml_path = "word_file/example_zip/word/fontTable.xml"
paragraphs, tables, xml_file, font_xml_file = read_word_document(input_file_path, input_file_xml_path,
                                                                 input_font_xml_path)

# 创建一个新的 Document 对象
doc = Document()

# 添加标题
doc.add_heading('My Document Title', level=1)

# 复制字符
copy_characters(paragraphs, doc, xml_file, font_xml_file)

# 添加带有项目符号的列表
doc.add_paragraph('This is a bulleted list:', style='ListBullet')
doc.add_paragraph('Item 1')
doc.add_paragraph('Item 2')

# 保存.docx文件
output_file_path = 'word_file/write_file.docx'
doc.save(output_file_path)

print('Document created successfully!')
