from docx import Document

# 创建一个新的Document对象
doc = Document()

# 添加标题
doc.add_heading('My Document Title', level=1)

# 添加段落
doc.add_paragraph('Hello, this is a paragraph.')

# 添加带有项目符号的列表
doc.add_paragraph('This is a bulleted list:', style='ListBullet')
doc.add_paragraph('Item 1')
doc.add_paragraph('Item 2')

# 保存.docx文件
doc.save('word_file/write_file.docx')

print('Document created successfully!')
